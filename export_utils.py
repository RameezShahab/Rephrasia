"""
export_utils.py — Export paraphrase / translation results to PDF and DOCX.

PDF export uses reportlab with proper reportlab color types.
DOCX export uses python-docx.
Both write to static/exports/ and return a relative URL path.
"""

import logging
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors as rl_colors           # ← reportlab colours
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

logger = logging.getLogger(__name__)

EXPORT_DIR = Path("static/exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def export_to_docx(original_text: str, results: list, result_type: str = "Paraphrased") -> str:
    """
    Export results to DOCX format.

    Args:
        original_text: The original input text.
        results:       List of result strings (paraphrases, translations, etc.).
        result_type:   Label shown in the document heading (e.g. 'Paraphrased').

    Returns:
        Relative URL path to the generated file (e.g. 'exports/abc123.docx').

    Raises:
        RuntimeError: On any docx generation failure.
    """
    try:
        filename = f"{uuid.uuid4().hex}.docx"
        filepath = EXPORT_DIR / filename

        doc = Document()

        doc.add_heading(f"{result_type} Text Results", level=0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        doc.add_heading("Original Text:", level=1)
        doc.add_paragraph(original_text)
        doc.add_paragraph("")

        doc.add_heading(f"{result_type} Versions:", level=1)
        if isinstance(results, list):
            for idx, result in enumerate(results, 1):
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. {result}")
                run.font.size = Pt(11)
                doc.add_paragraph("")
        else:
            doc.add_paragraph(str(results))

        doc.save(str(filepath))
        logger.info("DOCX exported: %s", filename)
        return f"exports/{filename}"

    except Exception as exc:
        logger.exception("DOCX export failed.")
        raise RuntimeError(f"DOCX export failed: {exc}") from exc


# ── PDF ───────────────────────────────────────────────────────────────────────

def export_to_pdf(original_text: str, results: list, result_type: str = "Paraphrased") -> str:
    """
    Export results to PDF format.

    Args:
        original_text: The original input text.
        results:       List of result strings.
        result_type:   Label shown in the document heading.

    Returns:
        Relative URL path to the generated file (e.g. 'exports/abc123.pdf').

    Raises:
        RuntimeError: On any PDF generation failure.
    """
    try:
        filename = f"{uuid.uuid4().hex}.pdf"
        filepath = EXPORT_DIR / filename

        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # ── Custom styles using reportlab colours (not python-docx RGBColor) ──
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=rl_colors.HexColor("#000080"),   # navy blue
            spaceAfter=12,
        )

        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=rl_colors.black,
            spaceAfter=6,
        )

        story.append(Paragraph(f"{result_type} Text Results", title_style))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            styles["Normal"],
        ))
        story.append(Spacer(1, 0.3 * inch))

        story.append(Paragraph("Original Text:", heading_style))
        story.append(Paragraph(original_text, styles["Normal"]))
        story.append(Spacer(1, 0.3 * inch))

        story.append(Paragraph(f"{result_type} Versions:", heading_style))
        story.append(Spacer(1, 0.1 * inch))

        if isinstance(results, list):
            for idx, result in enumerate(results, 1):
                story.append(Paragraph(f"{idx}. {result}", styles["Normal"]))
                story.append(Spacer(1, 0.1 * inch))
        else:
            story.append(Paragraph(str(results), styles["Normal"]))

        doc.build(story)
        logger.info("PDF exported: %s", filename)
        return f"exports/{filename}"

    except Exception as exc:
        logger.exception("PDF export failed.")
        raise RuntimeError(f"PDF export failed: {exc}") from exc


# ── Cleanup ───────────────────────────────────────────────────────────────────

def cleanup_old_exports(max_age_hours: int = 24) -> int:
    """
    Remove export files older than *max_age_hours*.

    Returns:
        Number of files deleted.
    """
    import time
    current_time = time.time()
    deleted = 0
    for export_file in EXPORT_DIR.glob("*.*"):
        age_hours = (current_time - export_file.stat().st_mtime) / 3600
        if age_hours > max_age_hours:
            export_file.unlink()
            deleted += 1
    if deleted:
        logger.info("Cleaned up %d old export file(s).", deleted)
    return deleted
